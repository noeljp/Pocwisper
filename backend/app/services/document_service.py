from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import os


class DocumentService:
    @staticmethod
    def create_docx(title: str, date: datetime, content: str, output_path: str) -> str:
        """Create a DOCX document from processed text"""
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"Date: {date.strftime('%d/%m/%Y')}").bold = True
        
        # Add separator
        doc.add_paragraph("_" * 50)
        
        # Add content
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
        
        # Save document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path


# Global instance
document_service = DocumentService()
